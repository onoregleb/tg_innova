from dotenv import load_dotenv
import os
from langchain_core.documents import Document
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.chat_models.gigachat import GigaChat
from langchain.chains import create_retrieval_chain

load_dotenv('.env')

# Ключ для GigaChat
giga_key = os.environ.get("SB_AUTH_DATA")


def load_docx(file_path: str) -> str:
    """
    Загружает текст из docx файла.

    Args:
        file_path (str): Путь к файлу docx.

    Returns:
        str: Текст из файла.
    """
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


def load_all_docx_from_folder(folder_path: str):
    """
    Загружает текст из всех docx файлов в указанной папке.

    Args:
        folder_path (str): Путь к папке.

    Returns:
        List[Document]: Список документов.
    """
    docs = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            file_path = os.path.join(folder_path, filename)
            document_text = load_docx(file_path)
            docs.append(Document(page_content=document_text))
    return docs


# Путь к папке с файлами .docx
folder_path = "docs"
docs = load_all_docx_from_folder(folder_path)

# Разделение текстов на части
text_splitter = RecursiveCharacterTextSplitter(chunk_size=100, chunk_overlap=20)
split_docs = text_splitter.split_documents(docs)

# Инициализация модели для создания эмбеддингов
model_name = "cointegrated/LaBSE-en-ru"
model_kwargs = {'device': 'cpu', 'trust_remote_code': True}  # Добавили trust_remote_code
encode_kwargs = {'normalize_embeddings': True}  # Для E5 важно нормализовать эмбеддинги
embedding = HuggingFaceEmbeddings(model_name=model_name, model_kwargs=model_kwargs, encode_kwargs=encode_kwargs)

# Проверяем, существует ли индекс
index_path = "faiss_index"

if os.path.exists(index_path):
    vector_store = FAISS.load_local(index_path, embedding, allow_dangerous_deserialization=True)
else:
    vector_store = FAISS.from_documents(split_docs, embedding)
    vector_store.save_local(index_path)

vector_store.save_local("faiss_index")

embedding_retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 4})


llm = GigaChat(credentials=giga_key, model='GigaChat', verify_ssl_certs=False, profanity_check=False)

prompt = ChatPromptTemplate.from_template('''Ты - AI-ассистент в строительной сфере. 
Ответь на вопрос пользователя, используя только информацию из контекста. 
Используй свои знания строительных терминов при ответе.
Если в контексте нет ответа, напиши: "Не могу ответить на Ваш вопрос. Попробуйте сформулировать иначе.". 
Не делай предположений.

Вопрос: {input}
Ответ: Вопрос: {input}
{context}''')

document_chain = create_stuff_documents_chain(llm=llm, prompt=prompt)

retrieval_chain = create_retrieval_chain(
    embedding_retriever,
    document_chain
)